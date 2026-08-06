from pathlib import Path

from docx import Document
from loguru import logger

from app.ingestion.document import DocumentPage, LoadedDocument


class DocxLoader:
    supported_extensions = {".docx"}

    def load(self, file_path: Path) -> LoadedDocument:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        if file_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"Unsupported DOCX file extension: {file_path.suffix}")

        logger.info("Loading DOCX document: {}", file_path)

        try:
            document = Document(file_path)
            paragraphs = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

        except Exception as error:
            logger.exception("Failed to load DOCX document: {}", file_path)
            raise RuntimeError(f"Failed to load DOCX document: {file_path}") from error

        text = "\n".join(paragraphs)

        page = DocumentPage(
            document_name=file_path.name,
            source_path=file_path,
            file_type="docx",
            text=text,
            page_number=None,
            unit_number=1,
            metadata={
                "loader": "DocxLoader",
                "paragraph_count": len(paragraphs),
            },
        )

        return LoadedDocument(
            document_name=file_path.name,
            source_path=file_path,
            file_type="docx",
            pages=[page],
        )